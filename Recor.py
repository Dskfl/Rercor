import sys
import os
import threading
import gc
import logging
import traceback
<<<<<<< HEAD
import cv2
import numpy as np
import ctypes

from pathlib import Path
from tkinter import Tk, filedialog, messagebox, Toplevel, Button, Label, Radiobutton, BooleanVar, StringVar, Frame, X
from tkinter import ttk
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from rembg import remove, new_session
from datetime import datetime

# Librerías para generación del documento en Word
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
=======
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65

# Configurar el archivo de registro (log) en el Escritorio
ruta_log = os.path.join(os.path.expanduser("~"), "Desktop", "Recor_Errores.log")
logging.basicConfig(filename=ruta_log, level=logging.ERROR, 
                    format='%(asctime)s - %(levelname)s - %(message)s')

def manejador_excepciones(exc_type, exc_value, exc_traceback):
    """Atrapa cualquier error fatal y lo guarda en el log en lugar de cerrar en silencio"""
    if issubclass(exc_type, KeyboardInterrupt):
        sys.__excepthook__(exc_type, exc_value, exc_traceback)
        return
    logging.error("Error crítico en la aplicación:", exc_info=(exc_type, exc_value, exc_traceback))

sys.excepthook = manejador_excepciones
<<<<<<< HEAD

=======
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
# Solución para el error 'NoneType' object has no attribute 'write'
if sys.stdout is None:
    sys.stdout = open(os.devnull, 'w')
if sys.stderr is None:
    sys.stderr = open(os.devnull, 'w')

global session_humanos
global session_general

<<<<<<< HEAD
=======
import cv2
import numpy as np
import ctypes

>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
# Esta línea le dice a Windows que tu ventana pertenece a una App única y no a Python genérico
myappid = 'estudio.recortador.carnet.1.0' 
try:
    ctypes.windll.shell32.SetCurrentProcessExplicitAppUserModelID(myappid)
except Exception:
    pass

<<<<<<< HEAD
=======
from pathlib import Path
from tkinter import Tk, filedialog, messagebox, Toplevel, Button, Label, Radiobutton, BooleanVar, StringVar, Frame, X
from tkinter import ttk
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from rembg import remove, new_session
from datetime import datetime

# OPTIMIZACIÓN: 1200 es suficiente para >300 DPI en tamaño cédula, aligera el PDF y el uso de RAM
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
OUTPUT_WIDTH = 1200 
PROPORCION_CEDULA = 0.6308411214953271
session_humanos = None
session_general = None

def mm_to_points(mm):
    return mm * 72 / 25.4

def ordenar_puntos(puntos):
    puntos = np.array(puntos, dtype=np.float32)
    suma = puntos.sum(axis=1)
    diferencia = np.diff(puntos, axis=1).reshape((-1))
    return np.array([puntos[np.argmin(suma)], puntos[np.argmin(diferencia)], 
                     puntos[np.argmax(suma)], puntos[np.argmax(diferencia)]], dtype=np.float32)

def detectar_con_ia_cedula(imagen):
    global session_general
    if session_general is None:
        session_general = new_session('u2net')
        
    alto, ancho = imagen.shape[:2]
    max_dim = 1200
    escala = 1
    if max(alto, ancho) > max_dim:
        escala = max_dim / max(alto, ancho)
        img_procesar = cv2.resize(imagen, (int(ancho * escala), int(alto * escala)), interpolation=cv2.INTER_AREA)
    else:
        img_procesar = imagen.copy()
        
    img_rgb = cv2.cvtColor(img_procesar, cv2.COLOR_BGR2RGB)
    img_sin_fondo = remove(img_rgb, session=session_general)
    mascara = img_sin_fondo[:, :, 3]
    contornos, _ = cv2.findContours(mascara, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    if not contornos:
        return
    else:
        contorno_max = max(contornos, key=cv2.contourArea)
        rectangulo = cv2.minAreaRect(contorno_max)
        return ordenar_puntos(cv2.boxPoints(rectangulo)) / escala

def recortar_cedula(imagen, esquinas):
    ancho_salida = OUTPUT_WIDTH
    alto_salida = int(ancho_salida * PROPORCION_CEDULA)
    destino = np.array([[0, 0], [ancho_salida - 1, 0], [ancho_salida - 1, alto_salida - 1], [0, alto_salida - 1]], dtype=np.float32)
    matriz = cv2.getPerspectiveTransform(esquinas, destino)
    return cv2.warpPerspective(imagen, matriz, (ancho_salida, alto_salida), flags=cv2.INTER_LANCZOS4)

def procesar_foto_cedula(ruta, salida, a_color=True):
    imagen = cv2.imread(str(ruta))
    if imagen is None:
        raise Exception(f'No se pudo cargar: {ruta}')
    else:
        esquinas = detectar_con_ia_cedula(imagen)
        if esquinas is None:
            alto, ancho = imagen.shape[:2]
            esquinas = np.array([[0, 0], [ancho, 0], [ancho, alto], [0, alto]], dtype=np.float32)
        resultado = recortar_cedula(imagen, esquinas)
        if not a_color:
            resultado = cv2.cvtColor(cv2.cvtColor(resultado, cv2.COLOR_BGR2GRAY), cv2.COLOR_GRAY2BGR)
        cv2.imwrite(str(salida), resultado, [cv2.IMWRITE_PNG_COMPRESSION, 1])
        
        del imagen
        del resultado
        gc.collect()

def crear_pdf_cedulas(pares_imagenes, archivo_pdf, usar_tamanio_grande):
    ancho_pagina, alto_pagina = letter
    margin = 36
    if usar_tamanio_grande:
        img_width = ancho_pagina - margin * 2
    else:
        img_width = mm_to_points(87.0)
    img_height = img_width * PROPORCION_CEDULA
    
    pdf = canvas.Canvas(str(archivo_pdf), pagesize=letter)
    x_pos = (ancho_pagina - img_width) / 2
    y_frente = alto_pagina - margin - img_height
    y_reverso = y_frente - 30 - img_height
    total = len(pares_imagenes)
    
    for i, (frente, reverso) in enumerate(pares_imagenes):
        pdf.drawImage(str(frente), x_pos, y_frente, width=img_width, height=img_height)
        pdf.drawImage(str(reverso), x_pos, y_reverso, width=img_width, height=img_height)
        if i < total - 1:
            pdf.showPage()
    pdf.save()

def procesar_foto_persona(ruta_in, ruta_out, modo_procesamiento):
    global session_humanos
    img = cv2.imread(ruta_in, cv2.IMREAD_UNCHANGED)
    if img is None:
        raise Exception(f'No se pudo cargar: {ruta_in}')
    else:
        alto, ancho = img.shape[:2]
<<<<<<< HEAD
        max_dim = 1024 
=======
        max_dim = 1024 # Escalado para evitar sobrecarga de memoria
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
        if max(alto, ancho) > max_dim:
            escala = max_dim / max(alto, ancho)
            img = cv2.resize(img, (int(ancho * escala), int(alto * escala)), interpolation=cv2.INTER_AREA)
            
        if modo_procesamiento == 'IA':
            if len(img.shape) == 2:
                img = cv2.cvtColor(img, cv2.COLOR_GRAY2BGR)
            elif len(img.shape) == 3 and img.shape[2] == 4:
                img = cv2.cvtColor(img, cv2.COLOR_BGRA2BGR)
                    
            img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
            
<<<<<<< HEAD
            if session_humanos is None:
                session_humanos = new_session('u2net_human_seg')
                
=======
            # OPTIMIZACIÓN IA: Modelo específico de humanos para no comerse partes del cuerpo
            if session_humanos is None:
                session_humanos = new_session('u2net_human_seg')
                
            # Aplicamos alpha_matting para perfeccionar los bordes (cabello fino)
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
            img_rgba = remove(
                img_rgb, 
                session=session_humanos, 
                post_process_mask=True,
                alpha_matting=True,
                alpha_matting_foreground_threshold=240,
                alpha_matting_background_threshold=10,
<<<<<<< HEAD
                alpha_matting_erode_size=5
            )
            
            img_bgra = cv2.cvtColor(img_rgba, cv2.COLOR_RGBA2BGRA)
            
            # --- ENCUADRE AJUSTADO Y MÁXIMO APROVECHAMIENTO ---
=======
                alpha_matting_erode_size=10
            )
            
            # ¡CORRECCIÓN DEL FILTRO VIOLETA AQUI! 
            # Volvemos a invertir los canales a BGR antes de pegarlo en el lienzo
            img_bgra = cv2.cvtColor(img_rgba, cv2.COLOR_RGBA2BGRA)
            
            # --- AUTO ENCUADRE Y CENTRADO ---
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
            mascara = img_bgra[:, :, 3]
            coordenadas = cv2.findNonZero(mascara)
            
            if coordenadas is not None:
<<<<<<< HEAD
                x, y, w, h = cv2.boundingRect(coordenadas)
                persona_recortada = img_bgra[y:y+h, x:x+w]
                
                # Reducimos el margen superior al 5% para que la persona llene el marco
                alto_lienzo = int(h * 1.05) 
                ancho_lienzo = int(alto_lienzo * (25 / 30))
                
                # Si los hombros son más anchos que el lienzo proporcional
                if w > ancho_lienzo: 
                    ancho_lienzo = int(w * 1.02)
                    alto_lienzo = int(ancho_lienzo * (30 / 25))
                
                lienzo = np.ones((alto_lienzo, ancho_lienzo, 3), dtype=np.uint8) * 255
                
                # Centrado horizontal y alineado al borde inferior
                x_offset = max(0, (ancho_lienzo - w) // 2)
                y_offset = max(0, alto_lienzo - h)
                
=======
                # Extraemos las coordenadas de la silueta real (Bounding Box)
                x, y, w, h = cv2.boundingRect(coordenadas)
                persona_recortada = img_bgra[y:y+h, x:x+w]
                
                # Proporción carnet (25x30 -> 5:6)
                alto_lienzo = int(h * 1.25) # 25% de margen extra en altura (cabeza/hombros)
                ancho_lienzo = int(alto_lienzo * (25 / 30))
                
                # Asegurar que la persona quepa horizontalmente
                if w > ancho_lienzo * 0.9: 
                    ancho_lienzo = int(w * 1.1)
                    alto_lienzo = int(ancho_lienzo * (30 / 25))
                
                # Crear lienzo blanco puro
                lienzo = np.ones((alto_lienzo, ancho_lienzo, 3), dtype=np.uint8) * 255
                
                # Centrado horizontal y margen inferior del 5%
                x_offset = (ancho_lienzo - w) // 2
                margen_inferior = int(alto_lienzo * 0.05)
                y_offset = alto_lienzo - h - margen_inferior
                
                # Evitar valores negativos por si las proporciones son inusuales
                x_offset = max(0, x_offset)
                y_offset = max(0, y_offset)
                
                # Superponer persona recortada en el lienzo blanco usando el canal Alfa
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
                alpha_fg = persona_recortada[:, :, 3] / 255.0
                for c in range(3):
                    lienzo[y_offset:y_offset+h, x_offset:x_offset+w, c] = \
                        (alpha_fg * persona_recortada[:, :, c] + (1 - alpha_fg) * 255).astype(np.uint8)
                
<<<<<<< HEAD
                img = lienzo 
            else:
                img = cv2.cvtColor(img_bgra, cv2.COLOR_BGRA2BGR)
                
        else:
            if len(img.shape) == 3 and img.shape[2] == 3:
                pass 
=======
                img = lienzo # La imagen final ya centrada y sin fondo transparente, con colores correctos
            else:
                # Si la IA no detecta a nadie, regresamos la original en blanco
                img = cv2.cvtColor(img_bgra, cv2.COLOR_BGRA2BGR)
                
        else:
            # Modo Manual (Photoshop)
            if len(img.shape) == 3 and img.shape[2] == 3:
                pass # Ya está en formato correcto, se asume fondo blanco
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
            elif len(img.shape) == 3 and img.shape[2] == 4:
                alpha_channel = img[:, :, 3] / 255.0
                white_bg = np.ones_like(img[:, :, :3]) * 255
                for c in range(3):
                    white_bg[:, :, c] = img[:, :, c] * alpha_channel + white_bg[:, :, c] * (1 - alpha_channel)
                img = white_bg
            
        cv2.imwrite(str(ruta_out), img, [cv2.IMWRITE_PNG_COMPRESSION, 1])
<<<<<<< HEAD
        del img
        gc.collect()

def crear_docx_carnet_4x6(rutas_multiplicadas, archivo_docx):
    """
    Genera un documento Word (.docx) configurado exactamente a 4x6 pulgadas 
    con una cuadrícula de 3x4 (12 fotos) ajustadas a 2.5 x 3.0 cm cada una.
    """
    doc = Document()
    
    # Definir tamaño de hoja a 4x6 pulgadas (10.16 cm x 15.24 cm)
    section = doc.sections[0]
    section.page_width = Inches(4)
    section.page_height = Inches(6)
    
    # Márgenes ajustados para centrar la cuadrícula de fotos
    section.top_margin = Inches(0.25)
    section.bottom_margin = Inches(0.25)
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    
    table = doc.add_table(rows=4, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for idx, ruta in enumerate(rutas_multiplicadas[:12]):
        fila = idx // 3
        col = idx % 3
        
        cell = table.cell(fila, col)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = paragraph.add_run()
        run.add_picture(str(ruta), width=Cm(2.5), height=Cm(3.0))
        
    archivo_salida = str(archivo_docx).replace('.pdf', '.docx')
    doc.save(archivo_salida)
    return archivo_salida
=======
        
        # Limpieza de memoria
        del img
        gc.collect()
def crear_pdf_carnet_4x6(rutas_multiplicadas, archivo_pdf):
    # --- Configuración del tamaño del Papel (4 x 6 pulgadas) ---
    ancho_pulgadas = 4.0
    alto_pulgadas = 6.0
    
    # ReportLab trabaja en "puntos" (1 pulgada = 72 puntos)
    ancho_papel = ancho_pulgadas * 72 
    alto_papel = alto_pulgadas * 72   
    
    pdf = canvas.Canvas(str(archivo_pdf), pagesize=(ancho_papel, alto_papel))
    
    # --- Configuración del tamaño de la Foto Carnet (2.5 cm ancho x 3.0 cm largo) ---
    ancho_cm = 2.5
    alto_cm = 3.0
    
    # Convertimos los cm a milímetros (multiplicando por 10) para la función mm_to_points
    w_pts = mm_to_points(ancho_cm * 10)
    h_pts = mm_to_points(alto_cm * 10)
    
    # Cuadrícula: 3 columnas x 4 filas = 12 fotos por hoja
    cols = 3
    rows = 4
    grid_w = cols * w_pts
    grid_h = rows * h_pts
    
    # Calculamos el espaciado automático entre las fotos para que queden centradas
    espacio_x = (ancho_papel - grid_w) / (cols + 1)
    espacio_y = (alto_papel - grid_h) / (rows + 1)
    
    max_por_pagina = cols * rows # 12
    
    for i, img_path in enumerate(rutas_multiplicadas):
        # Crear una nueva página si ya se llenó la actual
        if i > 0 and i % max_por_pagina == 0:
            pdf.showPage()
            
        idx_en_pagina = i % max_por_pagina
        r = idx_en_pagina // cols
        c = idx_en_pagina % cols
        
        # Calcular coordenadas X y Y para la foto actual
        x = espacio_x + c * (w_pts + espacio_x)
        y = alto_papel - (espacio_y + h_pts) - r * (h_pts + espacio_y)
        
        # Dibujar la imagen forzando la medida exacta
        pdf.drawImage(str(img_path), x, y, width=w_pts, height=h_pts, preserveAspectRatio=True, anchor='c')
        
    pdf.save()
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65

class AppEstudioCompleta:
    def __init__(self, root):
        self.root = root
        self.root.title('Gestor de impresiones v2.5 ')
        self.root.geometry('500x680') 
        self.root.resizable(False, False)
        style = ttk.Style()
        style.theme_use('clam')
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(expand=True, fill='both')
        
        self.btn_recortar_cedulas = None
        self.btn_recortar_carnet = None
        self.lbl_estado = None
        
        self.tab_cedulas = Frame(self.notebook, bg='#f9f9f9')
        self.notebook.add(self.tab_cedulas, text='🪪 Cédulas')
        self.construir_tab_cedulas()
        
        self.tab_carnet = Frame(self.notebook, bg='#f9f9f9')
        self.notebook.add(self.tab_carnet, text='📸 Fotos Carnet')
        self.construir_tab_carnet()
        
        self.lbl_estado = Label(self.root, text="", font=('Arial', 10, 'italic'), fg='blue')
        self.lbl_estado.pack(side="bottom", pady=5)

    def construir_tab_cedulas(self):
        self.frentes = []
        self.reversos = []
        self.var_color = BooleanVar(value=True)
        self.var_grande = BooleanVar(value=True)
        
        Label(self.tab_cedulas, text='Procesador Avanzado de Cédulas', font=('Arial', 14, 'bold'), bg='#f9f9f9').pack(pady=12)
        frame_opt = Frame(self.tab_cedulas, bd=2, relief='groove', padx=10, pady=10)
        frame_opt.pack(fill=X, padx=20, pady=5)
        Radiobutton(frame_opt, text='A Color', variable=self.var_color, value=True).pack(anchor='w')
        Radiobutton(frame_opt, text='Escala de Grises', variable=self.var_color, value=False).pack(anchor='w')
        Radiobutton(frame_opt, text='Tamaño ampliada', variable=self.var_grande, value=True).pack(anchor='w')
        Radiobutton(frame_opt, text='Tamaño Original (Como copia)', variable=self.var_grande, value=False).pack(anchor='w')
        
        frame_files = Frame(self.tab_cedulas, padx=10, pady=5, bg='#f9f9f9')
        frame_files.pack(fill=X, padx=20, pady=5)
        Button(frame_files, text='1. Seleccionar Frentes', command=self.cargar_frentes, bg='#fff9c4', font=('Arial', 10)).pack(fill=X, pady=4)
        self.lbl_frentes = Label(frame_files, text='Frentes: 0', fg='gray', bg='#f9f9f9')
        self.lbl_frentes.pack(anchor='w')
        
        Button(frame_files, text='2. Seleccionar Reversos', command=self.cargar_reversos, bg='#fff9c4', font=('Arial', 10)).pack(fill=X, pady=4)
        self.lbl_reversos = Label(frame_files, text='Reversos: 0', fg='gray', bg='#f9f9f9')
        self.lbl_reversos.pack(anchor='w')
        
        self.btn_recortar_cedulas = Button(self.tab_cedulas, text='✂️Recortar', command=self.procesar_cedulas, bg='#2FB3D4', fg='white', font=('Arial', 11, 'bold'), height=2)
        self.btn_recortar_cedulas.pack(fill=X, padx=20, pady=10)

    def construir_tab_carnet(self):
        self.fotos_carnet = []
        self.var_modo = StringVar(value='IA')
        self.var_cantidad = StringVar(value='4') 
        
<<<<<<< HEAD
        Label(self.tab_carnet, text='Generador Word 4x6', font=('Arial', 14, 'bold'), bg='#f9f9f9').pack(pady=12)
=======
        Label(self.tab_carnet, text='Generador de Hojas 4x6', font=('Arial', 14, 'bold'), bg='#f9f9f9').pack(pady=12)
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
        
        frame_modo = Frame(self.tab_carnet, bd=2, relief='groove', padx=10, pady=10)
        frame_modo.pack(fill=X, padx=20, pady=5)
        
        Label(frame_modo, text='Modo de Procesamiento:', font=('Arial', 10, 'bold')).pack(anchor='w')
        Radiobutton(frame_modo, text='IA (Automático - Quitar fondo)', variable=self.var_modo, value='IA').pack(anchor='w')
        Radiobutton(frame_modo, text='Manual (Photoshop / Ya editado)', variable=self.var_modo, value='Manual').pack(anchor='w')
        
        Label(frame_modo, text='Cantidad por cliente:', font=('Arial', 10, 'bold')).pack(anchor='w', pady=(10, 0))
        frame_radios_cant = Frame(frame_modo)
        frame_radios_cant.pack(anchor='w')
        for val in ['1', '2', '4', '6', '12']:
            Radiobutton(frame_radios_cant, text=val, variable=self.var_cantidad, value=val).pack(side='left', padx=5)

        frame_files = Frame(self.tab_carnet, padx=10, pady=5, bg='#f9f9f9')
        frame_files.pack(fill=X, padx=20, pady=5)
        Button(frame_files, text='Seleccionar Fotos de Clientes', command=self.cargar_carnets, bg='#fff9c4', font=('Arial', 11)).pack(fill=X, pady=4)
        self.lbl_carnets = Label(frame_files, text='Clientes: 0', fg='gray', bg='#f9f9f9')
        self.lbl_carnets.pack(anchor='w')
        
        self.btn_recortar_carnet = Button(self.tab_carnet, text='👤Recortar Fotos', command=self.procesar_carnets, bg='#2FB3D4', fg='white', font=('Arial', 11, 'bold'), height=2)
        self.btn_recortar_carnet.pack(fill=X, padx=20, pady=10)

    def cargar_frentes(self):
        archivos = filedialog.askopenfilenames(title='Selecciona FRENTES')
        if archivos:
            self.frentes = list(archivos)
            self.lbl_frentes.config(text=f'Frentes: {len(self.frentes)}', fg='green')

    def cargar_reversos(self):
        archivos = filedialog.askopenfilenames(title='Selecciona REVERSOS')
        if archivos:
            self.reversos = list(archivos)
            self.lbl_reversos.config(text=f'Reversos: {len(self.reversos)}', fg='green')

    def cargar_carnets(self):
        archivos = filedialog.askopenfilenames(title='Selecciona Fotos (PNG/JPG)')
        if archivos:
            self.fotos_carnet = list(archivos)
            self.lbl_carnets.config(text=f'Clientes: {len(self.fotos_carnet)}', fg='green')

    def procesar_cedulas(self):
        if not self.frentes or not self.reversos or len(self.frentes) != len(self.reversos):
            messagebox.showerror('Error', 'Revisa la selección de archivos. Debe haber igual cantidad de frentes que reversos.')
            return
        self.iniciar_proceso_hilo('cedulas')

    def procesar_carnets(self):
        if not self.fotos_carnet:
            messagebox.showerror('Error', 'Selecciona al menos una foto.')
            return
        self.iniciar_proceso_hilo('carnet')

    def iniciar_proceso_hilo(self, tipo):
        self.btn_recortar_cedulas.config(state='disabled')
        self.btn_recortar_carnet.config(state='disabled')
        self.lbl_estado.config(text="Procesando imágenes, por favor espera...", fg="blue")
        
        hilo = threading.Thread(target=self.ejecutar_generacion, args=(tipo,), daemon=True)
        hilo.start()

    def ejecutar_generacion(self, tipo):
        try:
            carpeta = Path.home() / 'OneDrive' / 'Desktop' / 'salidas_estudio'
            carpeta.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            
<<<<<<< HEAD
            archivo_final = None
            
=======
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
            if tipo == 'cedulas':
                pdf_salida = carpeta / f'cedulas_{timestamp}.pdf'
                pares_procesados = []
                for i, (f_ruta, r_ruta) in enumerate(zip(self.frentes, self.reversos)):
                    f_sal = carpeta / f'f_tmp_{i}.png'
                    r_sal = carpeta / f'r_tmp_{i}.png'
                    procesar_foto_cedula(f_ruta, f_sal, self.var_color.get())
                    procesar_foto_cedula(r_ruta, r_sal, self.var_color.get())
                    pares_procesados.append((f_sal, r_sal))
                crear_pdf_cedulas(pares_procesados, pdf_salida, self.var_grande.get())
<<<<<<< HEAD
                archivo_final = pdf_salida
                
            elif tipo == 'carnet':
                salida_base = carpeta / f'fotos_carnet_{timestamp}.docx'
=======
                
            elif tipo == 'carnet':
                pdf_salida = carpeta / f'fotos_carnet_{timestamp}.pdf'
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
                rutas_procesadas = []
                modo = self.var_modo.get()
                cantidad_por_cliente = int(self.var_cantidad.get()) 
                
                for i, ruta_img in enumerate(self.fotos_carnet):
                    img_sal = carpeta / f'c_tmp_{i}.png'
                    procesar_foto_persona(ruta_img, img_sal, modo)
                    rutas_procesadas.extend([img_sal] * cantidad_por_cliente)
                    
<<<<<<< HEAD
                # Generación en Word (.docx)
                archivo_final = crear_docx_carnet_4x6(rutas_procesadas, salida_base)
            
            gc.collect()
            
            if archivo_final:
                self.root.after(0, self.mostrar_exito, archivo_final)
            else:
                self.root.after(0, self.mostrar_error, "No se generó ningún archivo.")
=======
                crear_pdf_carnet_4x6(rutas_procesadas, pdf_salida)
            
            gc.collect()
            self.root.after(0, self.mostrar_exito, pdf_salida)
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
            
        except Exception as e:
            self.root.after(0, self.mostrar_error, str(e))

<<<<<<< HEAD
    def mostrar_exito(self, ruta_archivo):
=======
    def mostrar_exito(self, pdf_salida):
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65
        self.btn_recortar_cedulas.config(state='normal')
        self.btn_recortar_carnet.config(state='normal')
        self.lbl_estado.config(text="¡Proceso finalizado con éxito!", fg="green")
        
        sub = Toplevel(self.root)
<<<<<<< HEAD
        sub.title('Archivo Generado')
        sub.geometry('360x220')
        sub.resizable(False, False)
        Label(sub, text='¡Archivo generado con éxito!', font=('Arial', 11, 'bold'), pady=12).pack()
        Button(sub, text='🖨️ Abrir Documento Generado', command=lambda: [os.startfile(str(ruta_archivo)), sub.destroy()], bg='#4CAF50', fg='white', font=('Arial', 10), width=30, height=1).pack(pady=5)
        Button(sub, text='📂 Abrir Carpeta de Salidas', command=lambda: [os.startfile(str(Path(ruta_archivo).parent)), sub.destroy()], bg='#2196F3', fg='white', font=('Arial', 10), width=30, height=1).pack(pady=5)
=======
        sub.title('PDF Generado')
        sub.geometry('340x200')
        sub.resizable(False, False)
        Label(sub, text='¡PDF generado con éxito!', font=('Arial', 11, 'bold'), pady=12).pack()
        Button(sub, text='🖨️ Imprimir (Predeterminada)', command=lambda: [os.startfile(str(pdf_salida), 'print'), sub.destroy()], bg='#4CAF50', fg='white', font=('Arial', 10), width=28, height=1).pack(pady=5)
        Button(sub, text='📂 Abrir / Elegir Impresora (Manual)', command=lambda: [os.startfile(str(pdf_salida)), sub.destroy()], bg='#2196F3', fg='white', font=('Arial', 10), width=28, height=1).pack(pady=5)
>>>>>>> b63142d4fd71706c65d0ba673b6c0d4e9a058c65

    def mostrar_error(self, mensaje):
        self.btn_recortar_cedulas.config(state='normal')
        self.btn_recortar_carnet.config(state='normal')
        self.lbl_estado.config(text="Ocurrió un error.", fg="red")
        messagebox.showerror('Error', f'Detalles del error:\n{mensaje}')

def obtener_ruta_recurso(ruta_relativa):
    try:
        ruta_base = sys._MEIPASS
    except Exception:
        ruta_base = os.path.abspath(".")
    return os.path.join(ruta_base, ruta_relativa)

if __name__ == '__main__':
    root = Tk()
    try:
        ruta_icono = obtener_ruta_recurso('recor.ico')
        root.iconbitmap(ruta_icono)
    except Exception as e:
        print("No se pudo cargar el icono:", e)
        
    app = AppEstudioCompleta(root)
    root.mainloop()